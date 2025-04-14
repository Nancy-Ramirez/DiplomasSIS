# memorandum.py
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_memorandum(datos, archivo_masivo=None):
    # Obtener datos del formulario
    tipo_gestion = datos.get("tipoGestion")
    para = datos.get("para", "")
    de_ = datos.get("de", "")
    asunto = datos.get("asunto", "")
    fecha = datos.get("fecha", "")
    contenido = datos.get("contenido", "")
    tipo_letra = datos.get("letra", "")
    tamaño_letra = datos.get("tamañoLetra", "")
    alineacion = datos.get("alineacion", "")

    # Validar datos críticos
    try:
        tamaño_letra = int(tamaño_letra)
    except ValueError:
        tamaño_letra = 12  # Valor de respaldo por si viene mal

    alineacion_dict = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }
    alineacion_parrafo = alineacion_dict.get(alineacion, WD_ALIGN_PARAGRAPH.LEFT)

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Encabezado centrado
    header = doc.add_paragraph("MEMORÁNDUM")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_header = header.runs[0]
    run_header.bold = True
    run_header.font.size = Pt(tamaño_letra + 2)
    run_header.font.name = tipo_letra

    doc.add_paragraph("")  # Espacio

    # Datos básicos con estilo
    for label, value in [("Para", para), ("De", de_), ("Asunto", asunto), ("Fecha", fecha)]:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: {value}")
        run.bold = True
        run.font.name = tipo_letra
        run.font.size = Pt(tamaño_letra)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("")  # Espacio

    # Contenido justificado o según alineación
    cuerpo = doc.add_paragraph()
    run_cuerpo = cuerpo.add_run(contenido)
    run_cuerpo.font.name = tipo_letra
    run_cuerpo.font.size = Pt(tamaño_letra)
    cuerpo.alignment = alineacion_parrafo

    # Firma
    firma = doc.add_paragraph("\n\nAtentamente,\n\n__________________________")
    firma.runs[0].font.name = tipo_letra
    firma.runs[0].font.size = Pt(tamaño_letra)

    # Guardar
    nombre_archivo = f"memorandum_{para.replace(' ', '_')}.docx" if para else "memorandum.docx"
    output_path = os.path.join("uploads", nombre_archivo)
    doc.save(output_path)

    return output_path
