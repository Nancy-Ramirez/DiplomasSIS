from flask import Flask, request, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from docx.shared import Inches
import os

app = Flask(__name__)

@app.route('/')
def home():
    return "¡Servidor Flask en ejecución!"

# Carpeta donde se guardarán los archivos
UPLOAD_FOLDER = "uploads"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

@app.route('/recibir_datos', methods=['POST'])
def recibir_datos():
    try:
        # Verificar si se está recibiendo FormData
        if not request.form:
            return jsonify({"error": "Formato no soportado. Usa multipart/form-data"}), 415

        # Capturar los datos del formulario
        letra = request.form.get("letra", "No especificado")
        tamañoLetra = request.form.get("tamañoLetra", "No especificado")
        tamañoPapel = request.form.get("tamañoPapel", "No especificado")
        alineacion = request.form.get("alineacion", "No especificado")
        fecha = request.form.get("fecha", "No especificado")
        destinatario = request.form.get("destinatario", "No especificado")
        cuerpo = request.form.get("cuerpo", "No especificado")
        cantidadFirmantes = request.form.get("cantidadFirmantes", "0")

        # Capturar los nombres de los firmantes
        firmantes = []
        for i in range(1, int(cantidadFirmantes) + 1):
            firmante = request.form.get(f"firmante{i}", "No especificado")
            firmantes.append(firmante)

        # Guardar la imagen si se envió
        membrete_path = None
        if "membrete" in request.files:
            membrete = request.files["membrete"]
            if membrete.filename != "":
                filename = secure_filename(membrete.filename)
                membrete_path = os.path.join(UPLOAD_FOLDER, filename)
                membrete.save(membrete_path)

        # Crear el documento Word
        doc = Document()
        doc.add_paragraph(f"Fecha: {fecha}\n")
        doc.add_paragraph(f"Destinatario: {destinatario}\n")
        doc.add_paragraph("Cuerpo:\n")
        doc.add_paragraph(cuerpo)
        doc.add_paragraph(f"\nAlineación: {alineacion}")
        doc.add_paragraph(f"Tamaño de Letra: {tamañoLetra}")
        doc.add_paragraph(f"Tipo de Letra: {letra}")
        doc.add_paragraph(f"Tamaño de Papel: {tamañoPapel}\n")

        # Insertar la imagen si existe
        if membrete_path:
            doc.add_paragraph("\nMembrete:")
            doc.add_picture(membrete_path, width=Inches(2.5))  # Ajustar el tamaño de la imagen

        # Agregar los firmantes
        doc.add_paragraph("\nFirmantes:")
        for firmante in firmantes:
            doc.add_paragraph(f"- {firmante}")

        # Guardar el documento
        doc_path = os.path.join(UPLOAD_FOLDER, "carta_generada.docx")
        doc.save(doc_path)

        return jsonify({"mensaje": "Documento generado con éxito", "archivo_guardado": doc_path})

    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/verificar_sirve')
def verificar_sirve():
    return "si sirve"
