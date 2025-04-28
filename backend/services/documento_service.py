# 📄 backend/services/documento_service.py

from docx import Document
from docx.shared import Pt, Inches
import os
from werkzeug.utils import secure_filename
from datetime import datetime

# 📂 Carpeta de documentos generados
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 📅 Formatear fecha tipo "26 de abril de 2025"
def formatear_fecha(fecha_iso):
    try:
        fecha_dt = datetime.strptime(fecha_iso, "%Y-%m-%d")
        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        return f"{fecha_dt.day} de {meses[fecha_dt.month - 1]} de {fecha_dt.year}"
    except:
        return fecha_iso

# 📝 Generar documento Word
def generar_documento(datos, archivo_membrete):
    try:
        letra = datos.get("letra", "Arial")
        tamañoLetra = int(datos.get("tamañoLetra", 12))
        tamañoPapel = datos.get("tamañoPapel", "Carta")
        alineacion = datos.get("alineacion", "left")
        fecha = formatear_fecha(datos.get("fecha", "No especificado"))
        destinatario = datos.get("destinatario", "No especificado")
        cuerpo = datos.get("cuerpo", "No especificado")
        asunto = datos.get("asunto", "")
        cantidadFirmantes = datos.get("cantidadFirmantes", "0")

        # Firmantes
        firmantes = [datos.get(f"firmante{i}", "No especificado") for i in range(1, int(cantidadFirmantes) + 1)]

        # Guardar membrete (si hay)
        membrete_path = None
        if archivo_membrete:
            filename = secure_filename(archivo_membrete.filename)
            membrete_path = os.path.join(UPLOAD_FOLDER, filename)
            if not os.path.exists(membrete_path):
                archivo_membrete.save(membrete_path)

        # Crear documento
        doc = Document()

        style = doc.styles['Normal']
        style.font.name = letra
        style.font.size = Pt(tamañoLetra)

        # Tamaño de papel
        section = doc.sections[0]
        if tamañoPapel == "Carta":
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
        elif tamañoPapel == "Oficio":
            section.page_height = Inches(13)
            section.page_width = Inches(8.5)
        elif tamañoPapel == "Tabloide":
            section.page_height = Inches(17)
            section.page_width = Inches(11)
        elif tamañoPapel == "Ejecutivo":
            section.page_height = Inches(10.5)
            section.page_width = Inches(7.25)

        # Agregar membrete
        if membrete_path:
            p_membrete = doc.add_paragraph()
            run = p_membrete.add_run()
            run.add_picture(membrete_path, width=Inches(1.0))
            p_membrete.alignment = 2  # A la derecha

        # Agregar fecha
        p_fecha = doc.add_paragraph(fecha)
        p_fecha.alignment = 2

        # Agregar asunto
        if asunto:
            doc.add_paragraph("\n")
            p_asunto = doc.add_paragraph(asunto)
            p_asunto.alignment = 1  # Centrado

        # Agregar destinatario
        doc.add_paragraph("\n")
        doc.add_paragraph(destinatario)

        # Agregar cuerpo
        doc.add_paragraph("\n")
        p_cuerpo = doc.add_paragraph(cuerpo)
        alineaciones = {"left": 0, "center": 1, "right": 2, "justify": 3}
        p_cuerpo.alignment = alineaciones.get(alineacion, 0)

        # Firmantes
        for firmante in firmantes:
            doc.add_paragraph("_" * 30)
            doc.add_paragraph(f"{firmante}\n")

        # Guardar documento
        nombre_base = f"carta_{destinatario.replace(' ', '_')[:30]}" if destinatario != "No especificado" else "carta"
        nombre_archivo = secure_filename(f"{nombre_base}.docx")
        doc_path = os.path.join(UPLOAD_FOLDER, nombre_archivo)
        doc.save(doc_path)

        return doc_path, None  # ✅ Path y sin error

    except Exception as e:
        return None, str(e)  # ❌ Error
