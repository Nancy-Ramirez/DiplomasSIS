import os
import pandas as pd
from flask import jsonify
from docx import Document
from docx.shared import Pt, Inches
from werkzeug.utils import secure_filename
from zipfile import ZipFile

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), '..', 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# 🔵 Función para formatear la fecha
def formatear_fecha(fecha_iso):
    try:
        from datetime import datetime
        fecha_dt = datetime.strptime(fecha_iso, "%Y-%m-%d")
        meses = [
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        return f"{fecha_dt.day} de {meses[fecha_dt.month - 1]} de {fecha_dt.year}"
    except:
        return fecha_iso

# 🔵 Función para generar un memorandum individual
def generar_memorandum(datos):
    try:
        letra = datos.get("letra", "Arial")
        tamañoLetra = int(datos.get("tamañoLetra", 12))
        tamañoPapel = datos.get("tamañoPapel", "Carta")
        alineacion = datos.get("alineacion", "left")

        para = datos.get("para", "N/D")
        de = datos.get("de", "N/D")
        asunto = datos.get("asunto", "N/D")
        contenido = datos.get("contenido", "")
        fecha = formatear_fecha(datos.get("fecha", ""))

        doc = Document()

        # Estilo global
        style = doc.styles['Normal']
        style.font.name = letra
        style.font.size = Pt(tamañoLetra)

        # Tamaño de papel
        section = doc.sections[0]
        if tamañoPapel == "Carta":
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
        elif tamañoPapel == "Oficio":
            section.page_height = Inches(13)
            section.page_width = Inches(8.5)
        elif tamañoPapel == "Tabloide":
            section.page_height = Inches(17)
            section.page_width = Inches(11)
        elif tamañoPapel == "Ejecutivo":
            section.page_height = Inches(10.5)
            section.page_width = Inches(7.25)

        if fecha:
            p_fecha = doc.add_paragraph(fecha)
            p_fecha.alignment = 2  # Derecha

        # Título
        doc.add_paragraph("\nMEMORANDUM\n").alignment = 1  # Centro

        # Cuerpo del memorandum
        doc.add_paragraph(f"Para: {para}")
        doc.add_paragraph(f"De: {de}")
        doc.add_paragraph(f"Asunto: {asunto}")

        doc.add_paragraph("\n")
        p_cuerpo = doc.add_paragraph(contenido)
        alineaciones = {"left": 0, "center": 1, "right": 2, "justify": 3}
        p_cuerpo.alignment = alineaciones.get(alineacion, 0)

        # Guardar
        nombre_base = f"memorandum_{para.replace(' ', '_')[:30]}" if para != "N/D" else "memorandum"
        nombre_archivo = secure_filename(f"{nombre_base}.docx")
        path = os.path.join(UPLOAD_FOLDER, nombre_archivo)
        doc.save(path)

        return path

    except Exception as e:
        print(f"Error generando memorandum: {str(e)}")
        return None

# 🔵 Servicio principal
def crear_memorandum_service(request):
    try:
        if not request.form:
            return jsonify({"error": "Formato no soportado. Usa multipart/form-data"}), 415

        datos_formulario = request.form.to_dict()
        tipo_gestion = datos_formulario.get("tipoGestion", "individual")

        if tipo_gestion == "masiva":
            archivo_masivo = request.files.get("archivoMasivo")
            if not archivo_masivo:
                return jsonify({"error": "Falta el archivo para carga masiva"}), 400

            # Leer sin encabezados
            if archivo_masivo.filename.endswith(".xlsx"):
                df = pd.read_excel(archivo_masivo, header=None)
            elif archivo_masivo.filename.endswith(".csv"):
                df = pd.read_csv(archivo_masivo, header=None)
            else:
                return jsonify({"error": "Formato de archivo no soportado"}), 400

            zip_path = os.path.join(UPLOAD_FOLDER, "memorandums.zip")
            with ZipFile(zip_path, "w") as zipf:
                for index, row in df.iterrows():
                    para_valor = str(row.iloc[0]) if not pd.isna(row.iloc[0]) else "N/D"
                    datos_individuales = {
                        "para": para_valor,
                        "de": datos_formulario.get("de", "N/D"),
                        "asunto": datos_formulario.get("asunto", "N/D"),
                        "contenido": datos_formulario.get("contenido", "N/D"),
                        "fecha": datos_formulario.get("fecha", "N/D"),
                        "letra": datos_formulario.get("letra", "Arial"),
                        "tamañoLetra": datos_formulario.get("tamañoLetra", "12"),
                        "tamañoPapel": datos_formulario.get("tamañoPapel", "Carta"),
                        "alineacion": datos_formulario.get("alineacion", "left")
                    }
                    path = generar_memorandum(datos_individuales)
                    if path:
                        zipf.write(path, arcname=os.path.basename(path))

            return zip_path, None

        else:
            # Individual
            path = generar_memorandum(datos_formulario)
            return path, None

    except Exception as e:
        print(f"Error procesando memorandum: {str(e)}")
        return None, str(e)
